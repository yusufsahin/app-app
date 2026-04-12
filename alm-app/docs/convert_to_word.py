import os
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "#00FF00", "space": "0"},
        left={"sz": 24, "val": "dashed", "color": "#0000FF", "space": "0"},
        right={"sz": 24, "val": "dashed", "color": "#0000FF", "space": "0"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_doc():
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(10)

    # --- HEADER ---
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.text = "Pamera - Stratejik Süreç ve Kalite Yönetimi"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- LOGO & TITLE ---
    logo_url = "https://provera.com.tr/wp-content/uploads/2019/07/cropped-proveralogo-2.png"
    logo_path = "provera_logo.png"
    try:
        r = requests.get(logo_url, stream=True)
        if r.status_code == 200:
            with open(logo_path, 'wb') as f:
                for chunk in r:
                    f.write(chunk)
            doc.add_picture(logo_path, width=Inches(1.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    except Exception as e:
        print(f"Logo fetch failed: {e}")

    # --- TITLE ---
    title = doc.add_heading('Yönetici Özeti', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- HERO SECTION ---
    doc.add_heading('Bir karar almak için kaç kişiyle konuşmanız gerekiyor?', level=1)
    
    hero_text = (
        "Son aldığınız kritik karardan önce kaç farklı kaynaktan bilgi toplamak zorunda kaldınız? "
        "E-postalar, farklı ekiplerden gelen raporlar, birini arayıp güncel durumu sormak — "
        "bu bilgi derleme süreci her seferinde yeniden başlıyor. "
        "Ekibiniz haftada kaç saatini iş yapmak yerine bilgi aramaya harcıyor?"
    )
    doc.add_paragraph(hero_text)

    # --- STATS TABLE ---
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    
    cells[0].text = "Haftada 4 saat\n(Yöneticinin koordinasyon süresi)"
    cells[1].text = "3–5 farklı uygulama\n(Aynı bilginin tutulduğu yerler)"
    cells[2].text = "%60 toplantı\n(Sadece bilgi toplama amaçlı)"

    for cell in cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph() # Spacer

    # --- CORE PILLARS ---
    doc.add_heading('Pamera Değer Önerisi', level=2)
    p = doc.add_paragraph()
    p.add_run('🏛️ Tekil Platform: ').bold = True
    p.add_run('Tüm süreçler tek çatı altında. Tek noktadan yönetişim.')
    
    p = doc.add_paragraph()
    p.add_run('🔗 Tek Gerçek Kaynak: ').bold = True
    p.add_run('Uçtan uca izlenebilirlik. Herkes aynı veriye bakar.')

    p = doc.add_paragraph()
    p.add_run('⚙️ Sürecinize Uyan Yapı: ').bold = True
    p.add_run('Platform sizi kalıba zorlamaz. İş yapış biçiminize göre şekillenir.')

    doc.add_paragraph() # Spacer

    # --- PROBLEMS ---
    doc.add_heading('Gerçek Hayattaki Sorunlar', level=2)
    
    probs = [
        ("Körleşen Karar Kalitesi", "Genel müdür 'bu proje nerede?' diye sorduğunda cevap ertesi güne kalıyor. Kararlar gerçek verilere değil, o anki bilgiye erişime göre alınıyor."),
        ("Gizli Maliyet Birikimi", "Çoklu sistem lisansları ve araçlar arası manuel veri girişi tek bütçe satırında görünmüyor ama her yıl büyüyor."),
        ("Denetim Hazırlık Krizi", "Denetim öncesinde 'kim onayladı, belge nerede?' soruları günlerce aranır. Ekip işi bırakıp geçmişi karıştırır.")
    ]
    
    for p_title, p_desc in probs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{p_title}: ").bold = True
        p.add_run(p_desc)

    doc.add_page_break()

    # --- SOLUTIONS / OUTCOMES ---
    doc.add_heading('Pamera Çözümleri ve Kazanımlar', level=2)
    
    solutions = [
        ("Planlama ve İş Takibi", "Mevcut: Toplantı ve Excel zinciri.\nPamera ile: Hedefler ve zaman çizelgeleri tek ekranda. Karar hızı artar."),
        ("Gereksinim ve Onay Takibi", "Mevcut: Kimin onayladığı belirsiz.\nPamera ile: Her talep ve onay zaman damgalı kayıt altında. Hesap verebilirlik artar."),
        ("Kalite Kontrol ve Denetim", "Mevcut: Denetim listeleri Word/E-postada.\nPamera ile: Kontrol adımları anlık raporlanır. Denetim süresi kısalır."),
        ("Teslimat İzlenebilirliği", "Mevcut: 'Bu iş bitti mi?' sorusu herkese sorulur.\nPamera ile: Gereksinimden teslimata zincir tek yerden okunur.")
    ]

    for s_title, s_desc in solutions:
        doc.add_heading(s_title, level=3)
        doc.add_paragraph(s_desc)

    # --- FOOTER ---
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "Provera Telekomünikasyon ve Bilgi Teknolojileri | info@provera.com.tr"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('executive-summary-pamera.docx')
    print("executive-summary-pamera.docx generated successfully.")

if __name__ == "__main__":
    create_doc()
